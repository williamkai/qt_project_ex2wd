import os
from docx import Document
from PyQt6.QtCore import Qt
from PyQt6.QtWidgets import (
    QWidget, QPushButton, QApplication, QLabel, QComboBox,
    QLineEdit, QHBoxLayout,  QMessageBox,QProgressBar
)
from core.base_windows import BaseFuncWindow
from core.conversion_utils import (
    read_data_auto,
    col_letter_to_index,
    prepare_assign_map,
    replace_placeholders,
    save_doc_with_name,
)

class ExcelToInvitationWindow(BaseFuncWindow):
    def __init__(self, title="試算表轉召請文功能"):
        super().__init__(title)

    def setup_ui(self):
        super().setup_ui()

        layout = self.layout()

        def add_labeled_input(label_text, placeholder):
            row = QHBoxLayout()
            label = QLabel(label_text)
            label.setFixedWidth(100)
            label.setAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
            row.addWidget(label)

            line_edit = QLineEdit()
            line_edit.setPlaceholderText(placeholder)
            line_edit.setFixedWidth(200)
            row.addWidget(line_edit)

            wrapper = QWidget()
            wrapper.setLayout(row)
            layout.addWidget(wrapper, alignment=Qt.AlignmentFlag.AlignCenter)

            return line_edit
        
        def add_labeled_combobox(label_text, options, default_index=0):
            row = QHBoxLayout()
            label = QLabel(label_text)
            label.setFixedWidth(160)
            label.setAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
            row.addWidget(label)

            combo = QComboBox()
            combo.addItems([str(opt) for opt in options])
            combo.setCurrentIndex(default_index)
            combo.setFixedWidth(100)
            row.addWidget(combo)

            wrapper = QWidget()
            wrapper.setLayout(row)
            self.layout().addWidget(wrapper, alignment=Qt.AlignmentFlag.AlignCenter)

            return combo

        self.input_required = add_labeled_input("必須欄位：", "請輸入欄位，例如: c,d")
        self.input_optional = add_labeled_input("選擇欄位：", "請輸入欄位，例如: e,f")
        # 在 setup_ui 中加入
        self.font_size_1 = add_labeled_combobox("8字以內字體大小：", [8,10,12,14,16,18,20,22,25,28,30], default_index=7)  # 預設22
        self.font_size_2 = add_labeled_combobox("9~20字字體大小：", [8,10,12,14,16,18,20,22,25,28,30], default_index=5)  # 預設18
        self.font_size_3 = add_labeled_combobox("21字以上字體大小：", [6,8,10,12,14,16,18,20], default_index=3)       # 預設12

        btn_run = QPushButton("執行轉換")
        btn_run.setFixedSize(200, 40)
        btn_run.clicked.connect(self.run_conversion)
        layout.addWidget(btn_run, alignment=Qt.AlignmentFlag.AlignCenter)
        # ✅ 新增進度條
        self.progress_bar = QProgressBar()
        self.progress_bar.setFixedWidth(300)
        self.progress_bar.setValue(0)
        layout.addWidget(self.progress_bar, alignment=Qt.AlignmentFlag.AlignCenter)

    def run_conversion(self):
        if not self.excel_path:
            QMessageBox.warning(self, "錯誤", "請選擇 Excel 檔案。")
            return
        if not self.word_path:
            QMessageBox.warning(self, "錯誤", "請選擇 Word 模板檔案。")
            return
        if not self.output_folder:
            QMessageBox.warning(self, "錯誤", "請選擇輸出資料夾。")
            return

        required_text = self.input_required.text().strip()
        if not required_text:
            QMessageBox.warning(self, "錯誤", "請輸入必須欄位。")
            return

        if not all(col.strip().isalpha() and len(col.strip()) == 1 for col in required_text.split(",")):
            QMessageBox.warning(self, "錯誤", "必須欄位必須為英文單一字母，以逗號分隔。")
            return

        required_cols = [c.strip().upper() for c in required_text.split(",") if c.strip()]
        optional_text = self.input_optional.text().strip()
        optional_cols = [c.strip().upper() for c in optional_text.split(",") if c.strip()] if optional_text else []

        try:
            df = read_data_auto(self.excel_path)
            df = df.fillna('')
            # 限制最多只取前 200 筆
            # df = df.head(200)
            total_rows = len(df)
            self.progress_bar.setMaximum(total_rows)

            max_col_index = len(df.columns) - 1
            for col in required_cols + optional_cols:
                idx = col_letter_to_index(col)
                if idx < 0 or idx > max_col_index:
                    raise ValueError(f"欄位 {col} 超出檔案欄位範圍。")

            # 解析字體大小規則
            font_size_rules = [
                (8, int(self.font_size_1.currentText())),
                (20, int(self.font_size_2.currentText())),
                (9999, int(self.font_size_3.currentText())),
            ]

            for idx, row in df.iterrows():
                self.progress_bar.setValue(idx + 1)
                QApplication.processEvents()

                if not all(row[col_letter_to_index(c)] for c in required_cols):
                    continue

                doc = Document(self.word_path)

                # 產生替換字典
                replacements = {
                    c.upper(): str(row[col_letter_to_index(c)])
                    for c in required_cols + optional_cols
                }

                # 使用封裝的工具函數
                assign_map = prepare_assign_map(replacements, doc)
                replace_placeholders(doc, assign_map,font_size_rules)

                c_val = str(row[col_letter_to_index(required_cols[0])])[:6]
                filename = f"{c_val}_召請文.docx"
                folder_name = os.path.join(self.output_folder, "Excel 轉 Word 召請文")
                save_doc_with_name(doc, folder_name, filename)

            QMessageBox.information(self, "成功", "轉換完成！")
            self.progress_bar.setValue(total_rows)

        except Exception as e:
            QMessageBox.critical(self, "錯誤", f"轉換失敗：{str(e)}")
            self.progress_bar.setValue(0)


    
    # def run_conversion(self):
    #     # 基本驗證
    #     if not self.excel_path:
    #         QMessageBox.warning(self, "錯誤", "請選擇 Excel 檔案。")
    #         return
    #     if not self.word_path:
    #         QMessageBox.warning(self, "錯誤", "請選擇 Word 模板檔案。")
    #         return
    #     if not self.output_folder:
    #         QMessageBox.warning(self, "錯誤", "請選擇輸出資料夾。")
    #         return

    #     required_text = self.input_required.text().strip()
    #     if not required_text:
    #         QMessageBox.warning(self, "錯誤", "請輸入必須欄位。")
    #         return

    #     # 檢查必須欄位格式：一個字母，以逗號分隔
    #     if not all(col.strip().isalpha() and len(col.strip()) == 1 for col in required_text.split(",")):
    #         QMessageBox.warning(self, "錯誤", "必須欄位必須為英文單一字母，以逗號分隔。")
    #         return

    #     required_cols = [c.strip().upper() for c in required_text.split(",") if c.strip()]
    #     optional_text = self.input_optional.text().strip()
    #     optional_cols = [c.strip().upper() for c in optional_text.split(",") if c.strip()] if optional_text else []
            
    #     try:
    #         df = read_data_auto(self.excel_path)
    #         df = df.fillna('')
    #         total_rows = len(df)
    #         self.progress_bar.setMaximum(total_rows)

    #         max_col_index = len(df.columns) - 1
    #         for col in required_cols + optional_cols:
    #             idx = col_letter_to_index(col)
    #             if idx < 0 or idx > max_col_index:
    #                 raise ValueError(f"欄位 {col} 超出檔案欄位範圍。")

    #         for idx, row in df.iterrows():
    #             self.progress_bar.setValue(idx + 1)
    #             QApplication.processEvents()  # 即時更新 UI

    #             if not all(row[col_letter_to_index(c)] for c in required_cols):
    #                 continue

    #             doc = Document(self.word_path)

    #             # 產生替換值的 mapping
    #             replacements = {}
    #             for c in required_cols + optional_cols:
    #                 replacements[c.upper()] = str(row[col_letter_to_index(c)])

    #             # 統計所有佔位符出現次數
    #             placeholder_counts = {}
    #             pattern_template = r"{{\s*%s\s*}}"
    #             for key in replacements:
    #                 pattern = re.compile(pattern_template % key, re.IGNORECASE)
    #                 count = 0
    #                 for para in doc.paragraphs:
    #                     for run in para.runs:
    #                         count += len(pattern.findall(run.text))
    #                 placeholder_counts[key] = count
    #                 # print(f"佔位符 '{key}' 出現次數: {count}")

    #             # 建立 assign_map，每個佔位符要填哪些值
    #             assign_map = {}
    #             for key, raw_val in replacements.items():
    #                 vals = [v.strip() for v in raw_val.split(',') if v.strip()] if raw_val else ['']
    #                 n = placeholder_counts.get(key, 0)
    #                 if n == 0:
    #                     assign_map[key] = []
    #                 else:
    #                     if len(vals) < n:
    #                         vals += [''] * (n - len(vals))
    #                     elif len(vals) > n:
    #                         vals = vals[:n-1] + [','.join(vals[n-1:])]
    #                     assign_map[key] = vals
    #                 # print(f"key='{key}' 的替換清單: {assign_map[key]}")

    #             # 替換文字：逐個 run 處理，不破壞格式
    #             pattern_all = re.compile(r"{{\s*([a-zA-Z])\s*}}")
    #             key_counter = {k: 0 for k in assign_map}

    #             for para in doc.paragraphs:
    #                 for run in para.runs:
    #                     matches = list(pattern_all.finditer(run.text))
    #                     if not matches:
    #                         continue

    #                     new_text = run.text
    #                     for match in matches:
    #                         key = match.group(1).upper()
    #                         val_list = assign_map.get(key, [])
    #                         idx = key_counter.get(key, 0)
    #                         replacement = val_list[idx] if idx < len(val_list) else ''
    #                         key_counter[key] = idx + 1

    #                         # 動態設定字體大小
    #                         if len(replacement) <= 20:
    #                             font_size = Pt(22)
    #                         elif len(replacement) <= 25:
    #                             font_size = Pt(16)
    #                         else:
    #                             font_size = Pt(12)

    #                         new_text = new_text.replace(match.group(0), replacement, 1)
    #                         # print(f"run 內替換：{{{{{key}}}}} → '{replacement}'，字體：{font_size.pt}pt")

    #                     run.text = new_text
    #                     run.font.size = font_size  # <--- 動態字體大小
    #                     run.font.name = '標楷體'
    #                     run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

    #             # 儲存
    #             c_val = str(row[col_letter_to_index(required_cols[0])])[:6]
    #             filename = f"{c_val}_召請文.docx"
    #             folder_name = os.path.join(self.output_folder, "Excel 轉 Word 召請文")
    #             os.makedirs(folder_name, exist_ok=True)
    #             output_path = os.path.join(folder_name, filename)
    #             doc.save(output_path)

    #         QMessageBox.information(self, "成功", "轉換完成！")
    #         self.progress_bar.setValue(total_rows)
    #     except Exception as e:
    #         QMessageBox.critical(self, "錯誤", f"轉換失敗：{str(e)}")
    #         self.progress_bar.setValue(0)
            