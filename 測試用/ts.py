from docx import Document
import os

def inspect_word(file_path):
    doc = Document(file_path)

    print("🔍 文件分析開始")
    print("=" * 40)

    # 檢查段落
    print(f"📄 總段落數：{len(doc.paragraphs)}")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name
        print(f"段落 {i+1}: 「{text}」 (樣式: {style})")

    print("\n📊 檢查表格：")
    print(f"總表格數量：{len(doc.tables)}")
    for idx, table in enumerate(doc.tables):
        row_count = len(table.rows)
        col_count = len(table.columns)
        print(f"\n表格 {idx+1}: {row_count} 行 x {col_count} 欄")

        for r_idx, row in enumerate(table.rows):
            row_text = [cell.text.strip() for cell in row.cells]
            print(f"  第 {r_idx+1} 行: {row_text}")

    print("=" * 40)
    print("✅ 分析完成")

if __name__ == "__main__":
    try:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        word_path = os.path.join(base_dir, "114拔薦文疏-本 空白_模板.docx")

        if not os.path.exists(word_path):
            raise FileNotFoundError("❌ 找不到 Word 檔案")

        inspect_word(word_path)

    except Exception as e:
        print("❌ 錯誤：", e)
