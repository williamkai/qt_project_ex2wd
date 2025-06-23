from PyQt6.QtCore import QObject, pyqtSignal, QRunnable
from docx import Document
from core.conversion_utils import (
    read_data_auto,
    prepare_assign_map,
    replace_placeholders,
    save_doc_with_name,
    map_all_placeholders,
    fill_data_to_table_v2,
    duplicate_table_and_insert
)
import os
from copy import deepcopy
class ExportWorkerSignals(QObject):
    finished = pyqtSignal(bool, str)
    progress = pyqtSignal(int)

class ExportWorker(QRunnable):
    def __init__(self, exporter, output_path):
        super().__init__()
        self.exporter = exporter
        self.output_path = output_path
        self.signals = ExportWorkerSignals()

    def run(self):
        success, result = self.exporter.export(
            self.output_path,
            progress_callback=self.signals.progress.emit
        )
        self.signals.finished.emit(success, result)


class WordExportWorkerSignals(QObject):
    finished = pyqtSignal(bool, str)
    progress = pyqtSignal(int)

class WordExportWorker(QRunnable):
    def __init__(
        self,
        excel_path,
        word_path,
        output_folder,
        required_cols,
        optional_cols,
        font_size_rules,
        limit_rows,
    ):
        super().__init__()
        self.signals = WordExportWorkerSignals()
        self.excel_path = excel_path
        self.word_path = word_path
        self.output_folder = output_folder
        self.required_cols = required_cols
        self.optional_cols = optional_cols
        self.font_size_rules = font_size_rules
        self.limit_rows = limit_rows

    def run(self):
        try:
            df = read_data_auto(self.excel_path).fillna('')
            if self.limit_rows and isinstance(self.limit_rows, int):
                df = df.head(self.limit_rows)

            col_letter_map = {chr(65+i): name for i, name in enumerate(df.columns)}

            for col in self.required_cols + self.optional_cols:
                if col not in col_letter_map:
                    raise ValueError(f"欄位 {col} 不存在於讀取的資料中。")

            # 預先計算應該會寫出的檔案數（有填寫必填欄位的筆數）
            valid_rows = [row for idx, row in df.iterrows() if all(str(row[col_letter_map[c]]) for c in self.required_cols)]
            total_valid = len(valid_rows)

            # 發送進度條最大值
            self.signals.progress.emit(0)  # 先重置，實際在UI設最大值，這行可以省略看你UI怎寫

            written_count = 0
            for row in valid_rows:
                doc = Document(self.word_path)
                replacements = {c: str(row[col_letter_map[c]]) for c in self.required_cols + self.optional_cols}
                assign_map = prepare_assign_map(replacements, doc)
                replace_placeholders(doc, assign_map, self.font_size_rules)

                c_val = str(row[col_letter_map[self.required_cols[0]]])[:11]
                filename = f"{c_val}_召請文.docx"
                folder_name = os.path.join(self.output_folder, "Excel 轉 Word 召請文")
                save_doc_with_name(doc, folder_name, filename)

                written_count += 1
                self.signals.progress.emit(written_count)

            self.signals.finished.emit(True, "轉換完成！")

        except Exception as e:
            self.signals.finished.emit(False, str(e))




class WordBatchExportWorkerSignals(QObject):
    finished = pyqtSignal(bool, str)  # (成功, 訊息)
    progress = pyqtSignal(int)  # 進度值

class WordBatchExportWorker(QRunnable):
    def __init__(
        self,
        excel_path,
        word_path,
        output_folder,
        required_cols,
        optional_cols,
        limit_rows,
        is_closing_getter,  # 傳入一個 func 取得 is_closing 狀態
        get_font_size_func,
    ):
        super().__init__()
        self.signals = WordBatchExportWorkerSignals()
        self.excel_path = excel_path
        self.word_path = word_path
        self.output_folder = output_folder
        self.required_cols = required_cols
        self.optional_cols = optional_cols
        self.limit_rows = limit_rows
        self.is_closing_getter = is_closing_getter
        self.get_font_size_func = get_font_size_func

    def run(self):
        try:
            df = read_data_auto(self.excel_path).fillna('')
            if self.limit_rows and isinstance(self.limit_rows, int):
                df = df.head(self.limit_rows)

            col_letter_map = {chr(65+i): name for i, name in enumerate(df.columns)}

            for col in self.required_cols + self.optional_cols:
                if col not in col_letter_map:
                    raise ValueError(f"欄位 {col} 不存在於讀取的資料中。")

            total_rows = len(df)
            doc = Document(self.word_path)
            if len(doc.tables) == 0:
                raise ValueError("找不到 Word 表格")

            table = doc.tables[0]
            clean_template_table = deepcopy(table)
            placeholder_map, start_col = map_all_placeholders(table)
            col_count = start_col + 1

            all_data = []
            for idx, row in df.iterrows():
                if self.is_closing_getter():
                    self.signals.progress.emit(0)
                    self.signals.finished.emit(False, "使用者中止轉換")
                    return

                if not all(str(row[col_letter_map[c]]) for c in self.required_cols):
                    continue

                replacements = {c: str(row[col_letter_map[c]]) for c in self.required_cols + self.optional_cols}
                all_data.append(replacements)

            font_size_func = self.get_font_size_func()

            batch_start = 0
            current_table = table
            while batch_start < len(all_data):
                remaining = len(all_data) - batch_start
                fill_count = min(col_count, remaining)

                current_batch = all_data[batch_start:batch_start + fill_count]
                written = fill_data_to_table_v2(
                    current_table,
                    placeholder_map,
                    current_batch,
                    start_col,
                    font_size_func=font_size_func
                )

                if written == 0:
                    raise RuntimeError(
                        f"填入資料失敗：\n\n"
                        f"可能模板錯誤或欄位不匹配。\n"
                        f"目前要填入的資料：\n{current_batch[:1]}"
                    )

                batch_start += written
                self.signals.progress.emit(batch_start)

                if batch_start < len(all_data):
                    current_table = duplicate_table_and_insert(doc, clean_template_table)

            out_path = os.path.join(self.output_folder, "牌位文疏批次生成.docx")
            doc.save(out_path)

            self.signals.finished.emit(True, "轉換完成！")

        except Exception as e:
            self.signals.finished.emit(False, f"轉換失敗：{str(e)}")