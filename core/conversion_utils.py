# ✅ conversion_utils.py
import os
import re
import pandas as pd
from docx.shared import Pt
from docx.oxml.ns import qn
import os
from docx import Document
from docx.enum.text import WD_BREAK
from copy import deepcopy
import pandas as pd
import re
from docx.shared import Pt
from docx.oxml.ns import qn
__all__ = [
    "read_data_auto",
    "col_letter_to_index",
    "prepare_assign_map",
    "replace_placeholders",
    "save_doc_with_name",
    "write_text_to_cell",
    "find_placeholders",
    "fill_data_to_table",
    "duplicate_table_and_insert",
    "extract_placeholders_from_cell",
    "map_all_placeholders",
    "fill_data_to_table_v2",
    "get_dynamic_font_size_func"
    ]

def read_data_auto(filepath):
    try:
        with open(filepath, 'rb') as f:
            head = f.read(1024)
    except Exception as e:
        raise ValueError(f"無法開啟檔案: {e}")

    try:
        if head.startswith(b'\xD0\xCF\x11\xE0'):
            return pd.read_excel(filepath, dtype=str, engine='xlrd')
        elif head.startswith(b'PK\x03\x04'):
            return pd.read_excel(filepath, dtype=str, engine='openpyxl')
        elif b'<html' in head.lower() or b'<!doctype' in head.lower():
            return pd.read_html(filepath, encoding='utf-8')[0]
        else:
            return pd.read_csv(filepath, dtype=str, encoding='utf-8')
    except Exception as e:
        raise ValueError(f"讀取檔案失敗: {e}")

def col_letter_to_index(letter):
    return ord(letter.upper()) - ord('A')


def prepare_assign_map(replacements, doc):
    assign_map = {}
    for key, raw_val in replacements.items():
        val = raw_val.strip() if raw_val else ''
        assign_map[key] = val
    return assign_map


def replace_placeholders(doc, assign_map, font_size_rules=None):
    if font_size_rules is None:
        font_size_rules = [(8, 22), (20, 18), (9999, 12)]

    pattern = re.compile(r"{{\s*([a-zA-Z]+)\s*}}")

    for para in doc.paragraphs:
        runs = para.runs
        if not runs:
            continue

        # 合併段落所有 run 的文字，並記錄每個 run 的 (start_idx, end_idx, run_index)
        full_text = ''
        run_positions = []  # (start_idx, end_idx, run_idx)
        idx = 0
        for i, run in enumerate(runs):
            text = run.text
            start_idx = idx
            end_idx = start_idx + len(text)
            run_positions.append((start_idx, end_idx, i))
            full_text += text
            idx = end_idx

        # 找到所有佔位符（從後往前處理避免索引錯亂）
        for match in reversed(list(pattern.finditer(full_text))):
            key = match.group(1).upper()
            replacement = assign_map.get(key, '')

            start, end = match.start(), match.end()

            # 找出被佔位符覆蓋的 runs
            affected = [r for r in run_positions if not (r[1] <= start or r[0] >= end)]
            if not affected:
                continue

            first_run_idx = affected[0][2]

            # 先將第一個 run 的文字替換成替換文字
            runs[first_run_idx].text = replacement

            # 其他被佔位符覆蓋的 run 文字清空
            for _, _, r_idx in affected[1:]:
                runs[r_idx].text = ''

            # 設定字體大小及字型
            length = len(replacement)
            font_size = Pt(12)
            for max_len, size_pt in font_size_rules:
                if length <= max_len:
                    font_size = Pt(size_pt)
                    break

            run = runs[first_run_idx]
            run.font.size = font_size
            run.font.name = '標楷體'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

def save_doc_with_name(doc, folder, filename):
    os.makedirs(folder, exist_ok=True)
    output_path = os.path.join(folder, filename)
    doc.save(output_path)
    return output_path
##############################################################

def extract_placeholders_from_cell(cell):
    pattern = re.compile(r"{{\s*([a-zA-Z]+)\s*}}")
    full_text = "".join(run.text for para in cell.paragraphs for run in para.runs)
    return pattern.findall(full_text)


def map_all_placeholders(table):
    """
    建立 col: [(row, [keys])] 的 dict，表示此欄位的哪些 row 含有哪些 {{key}}
    並回傳最大欄位（右側起點）作為 start_col。
    """
    mapping = {}
    max_col = 0
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            keys = extract_placeholders_from_cell(cell)
            if keys:
                mapping.setdefault(c_idx, []).append((r_idx, [k.upper() for k in keys]))
                max_col = max(max_col, c_idx)
    return mapping, max_col


def write_text_to_cell(cell, text, font_size=12):
    cell.text = str(text)
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = para.runs[0] if para.runs else para.add_run()
    run.font.size = Pt(font_size)
    run.font.name = '標楷體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')


def fill_data_to_table_v2(table, placeholder_map, data_batch, start_col, font_size_func=None):
    """
    根據 placeholder_map 的最右一欄做為模板欄（e.g., 14），
    從右到左將資料寫入每一欄（最大支援 15 筆）。
    font_size_func: 一個函式，根據每筆資料內容長度決定字體大小。
    """
    for i, data in enumerate(data_batch):
        col = start_col - i
        if col < 0:
            break
        template_col = start_col
        if template_col not in placeholder_map:
            continue
        for row, keys in placeholder_map[template_col]:
            val = "\n".join(data.get(k, '') for k in keys)
            font_size = font_size_func(val) if font_size_func else 12
            write_text_to_cell(table.cell(row, col), val, font_size=font_size)
    return len(data_batch)

def duplicate_table_and_insert(doc, table):
    new_table = deepcopy(table)

    # 新增段落並加上分頁符號
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)

    # 將表格插入在分頁符號之後
    # 找到 paragraph 的 XML 元素
    p_element = paragraph._element

    # 插入在該段落之後
    p_element.addnext(new_table._element)

    return new_table


def get_dynamic_font_size_func(size_1, size_2, size_3):
    def font_size_func(text):
        length = len(text.replace("\n", ""))
        if length <= 6:
            return size_1
        elif length <= 15:
            return size_2
        else:
            return size_3
    return font_size_func
