import os,re
import pandas as pd
from docx import Document
from PyQt6.QtCore import Qt, QSize
from PyQt6.QtGui import QFontMetrics
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QPushButton, QApplication, QLabel, QFileDialog,
    QSizePolicy, QLineEdit, QHBoxLayout, QFormLayout, QMessageBox,QProgressBar
)
from docx.shared import Pt
from docx.oxml.ns import qn

class BaseFuncWindow(QWidget):
    def __init__(self, title):
        super().__init__()
        self.setWindowTitle(title)
        self.setMinimumSize(600, 400)

        self.excel_path = None
        self.word_path = None
        self.output_folder = None

        self.setup_ui()

    def setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)

        def add_labeled_button(label_text, button_text, click_func):
            container = QHBoxLayout()

            label = QLabel(label_text)
            label.setFixedWidth(100)
            label.setAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
            container.addWidget(label)

            button = QPushButton(button_text)
            button.setFixedSize(200, 40)
            button.clicked.connect(click_func)
            container.addWidget(button)

            wrapper = QWidget()
            wrapper.setLayout(container)
            layout.addWidget(wrapper, alignment=Qt.AlignmentFlag.AlignCenter)

            return button

        self.btn_select_excel = add_labeled_button("Excel 檔案：", "選擇 Excel 檔案", self.select_excel_file)
        self.btn_select_word = add_labeled_button("Word 模板：", "選擇 Word 模板", self.select_word_file)
        self.btn_select_output = add_labeled_button("輸出資料夾：", "選擇輸出資料夾", self.select_output_folder)

        self.setLayout(layout)

    def select_excel_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "選擇 Excel 檔案", "", "Excel 檔案 (*.xls *.xlsx)")
        if path:
            self.excel_path = path
            filename = os.path.basename(path)
            self.btn_select_excel.setText(f"Excel: {filename}")

    def select_word_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "選擇 Word 模板", "", "Word 文件 (*.docx)")
        if path:
            self.word_path = path
            filename = os.path.basename(path)
            self.btn_select_word.setText(filename)

    def select_output_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "選擇輸出資料夾")
        if folder:
            self.output_folder = folder
            metrics = QFontMetrics(self.btn_select_output.font())
            max_width = 300
            elided_text = metrics.elidedText(folder, Qt.TextElideMode.ElideMiddle, max_width)
            self.btn_select_output.setText(elided_text)
            self.btn_select_output.setMaximumWidth(max_width)
            self.btn_select_output.setSizePolicy(QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Fixed)


class ExcelToInvitationWindow(BaseFuncWindow):
    def __init__(self, title="試算表轉召請文功能"):
        super().__init__(title)

    def setup_ui(self):
        super().setup_ui()

        layout = self.layout()

        def add_labeled_input(label_text, placeholder):
            row = QHBoxLayout()
            label = QLabel(label_text)
            label.setFixedWidth(100)
            label.setAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
            row.addWidget(label)

            line_edit = QLineEdit()
            line_edit.setPlaceholderText(placeholder)
            line_edit.setFixedWidth(200)
            row.addWidget(line_edit)

            wrapper = QWidget()
            wrapper.setLayout(row)
            layout.addWidget(wrapper, alignment=Qt.AlignmentFlag.AlignCenter)

            return line_edit

        self.input_required = add_labeled_input("必須欄位：", "請輸入欄位，例如: c,d")
        self.input_optional = add_labeled_input("選擇欄位：", "請輸入欄位，例如: e,f")

        btn_run = QPushButton("執行轉換")
        btn_run.setFixedSize(200, 40)
        btn_run.clicked.connect(self.run_conversion)
        layout.addWidget(btn_run, alignment=Qt.AlignmentFlag.AlignCenter)
        # ✅ 新增進度條
        self.progress_bar = QProgressBar()
        self.progress_bar.setFixedWidth(300)
        self.progress_bar.setValue(0)
        layout.addWidget(self.progress_bar, alignment=Qt.AlignmentFlag.AlignCenter)


    def run_conversion(self):
        # 基本驗證
        if not self.excel_path:
            QMessageBox.warning(self, "錯誤", "請選擇 Excel 檔案。")
            return
        if not self.word_path:
            QMessageBox.warning(self, "錯誤", "請選擇 Word 模板檔案。")
            return
        if not self.output_folder:
            QMessageBox.warning(self, "錯誤", "請選擇輸出資料夾。")
            return

        required_text = self.input_required.text().strip()
        if not required_text:
            QMessageBox.warning(self, "錯誤", "請輸入必須欄位。")
            return

        # 檢查必須欄位格式：一個字母，以逗號分隔
        if not all(col.strip().isalpha() and len(col.strip()) == 1 for col in required_text.split(",")):
            QMessageBox.warning(self, "錯誤", "必須欄位必須為英文單一字母，以逗號分隔。")
            return

        required_cols = [c.strip().upper() for c in required_text.split(",") if c.strip()]
        optional_text = self.input_optional.text().strip()
        optional_cols = [c.strip().upper() for c in optional_text.split(",") if c.strip()] if optional_text else []

        def read_data_auto(filepath):
            try:
                with open(filepath, 'rb') as f:
                    head = f.read(1024)
            except Exception as e:
                raise ValueError(f"無法開啟檔案: {e}")

            try:
                if head.startswith(b'\xD0\xCF\x11\xE0'):
                    # .xls 格式
                    try:
                        return pd.read_excel(filepath, dtype=str, engine='xlrd')
                    except Exception:
                        # 如果讀取xls失敗，嘗試用read_html fallback
                        dfs = pd.read_html(filepath, encoding='utf-8')
                        return dfs[0]
                elif head.startswith(b'PK\x03\x04'):
                    # .xlsx 格式
                    return pd.read_excel(filepath, dtype=str, engine='openpyxl')
                elif b'<html' in head.lower() or b'<!doctype' in head.lower():
                    dfs = pd.read_html(filepath, encoding='utf-8')
                    return dfs[0]
                else:
                    # 嘗試讀 CSV
                    try:
                        return pd.read_csv(filepath, dtype=str, encoding='utf-8')
                    except Exception:
                        raise ValueError("檔案格式不支援，請使用 .xls, .xlsx, CSV 或 HTML 格式。")
            except Exception as e:
                raise ValueError(f"讀取檔案失敗: {e}")
            
        try:
            df = read_data_auto(self.excel_path)
            df = df.fillna('')
            total_rows = len(df)
            self.progress_bar.setMaximum(total_rows)

            def col_letter_to_index(letter):
                return ord(letter.upper()) - ord('A')

            max_col_index = len(df.columns) - 1
            for col in required_cols + optional_cols:
                idx = col_letter_to_index(col)
                if idx < 0 or idx > max_col_index:
                    raise ValueError(f"欄位 {col} 超出檔案欄位範圍。")

            for idx, row in df.iterrows():
                self.progress_bar.setValue(idx + 1)
                QApplication.processEvents()  # 即時更新 UI

                if not all(row[col_letter_to_index(c)] for c in required_cols):
                    continue

                doc = Document(self.word_path)

                # 產生替換值的 mapping
                replacements = {}
                for c in required_cols + optional_cols:
                    replacements[c.upper()] = str(row[col_letter_to_index(c)])

                # 統計所有佔位符出現次數
                placeholder_counts = {}
                pattern_template = r"{{\s*%s\s*}}"
                for key in replacements:
                    pattern = re.compile(pattern_template % key, re.IGNORECASE)
                    count = 0
                    for para in doc.paragraphs:
                        for run in para.runs:
                            count += len(pattern.findall(run.text))
                    placeholder_counts[key] = count
                    # print(f"佔位符 '{key}' 出現次數: {count}")

                # 建立 assign_map，每個佔位符要填哪些值
                assign_map = {}
                for key, raw_val in replacements.items():
                    vals = [v.strip() for v in raw_val.split(',') if v.strip()] if raw_val else ['']
                    n = placeholder_counts.get(key, 0)
                    if n == 0:
                        assign_map[key] = []
                    else:
                        if len(vals) < n:
                            vals += [''] * (n - len(vals))
                        elif len(vals) > n:
                            vals = vals[:n-1] + [','.join(vals[n-1:])]
                        assign_map[key] = vals
                    # print(f"key='{key}' 的替換清單: {assign_map[key]}")

                # 替換文字：逐個 run 處理，不破壞格式
                pattern_all = re.compile(r"{{\s*([a-zA-Z])\s*}}")
                key_counter = {k: 0 for k in assign_map}

                for para in doc.paragraphs:
                    for run in para.runs:
                        matches = list(pattern_all.finditer(run.text))
                        if not matches:
                            continue

                        new_text = run.text
                        for match in matches:
                            key = match.group(1).upper()
                            val_list = assign_map.get(key, [])
                            idx = key_counter.get(key, 0)
                            replacement = val_list[idx] if idx < len(val_list) else ''
                            key_counter[key] = idx + 1

                            # 動態設定字體大小
                            if len(replacement) <= 20:
                                font_size = Pt(22)
                            elif len(replacement) <= 25:
                                font_size = Pt(16)
                            else:
                                font_size = Pt(12)

                            new_text = new_text.replace(match.group(0), replacement, 1)
                            # print(f"run 內替換：{{{{{key}}}}} → '{replacement}'，字體：{font_size.pt}pt")

                        run.text = new_text
                        run.font.size = font_size  # <--- 動態字體大小
                        run.font.name = '標楷體'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

                # 儲存
                c_val = str(row[col_letter_to_index(required_cols[0])])[:6]
                filename = f"{c_val}_召請文.docx"
                folder_name = os.path.join(self.output_folder, "Excel 轉 Word 召請文")
                os.makedirs(folder_name, exist_ok=True)
                output_path = os.path.join(folder_name, filename)
                doc.save(output_path)

            QMessageBox.information(self, "成功", "轉換完成！")
            self.progress_bar.setValue(total_rows)
        except Exception as e:
            QMessageBox.critical(self, "錯誤", f"轉換失敗：{str(e)}")
            self.progress_bar.setValue(0)
            
class InvitationTransferWindow(BaseFuncWindow):
    def __init__(self, title="拔薦文疏-本轉換功能"):
        super().__init__(title)

    def setup_ui(self):
        super().setup_ui()
        label_func = QLabel("拔薦文疏-本轉換功能（還沒寫）")
        self.layout().addWidget(label_func)


class ReceiptTransferWindow(BaseFuncWindow):
    def __init__(self, title="收據轉換功能"):
        super().__init__(title)

    def setup_ui(self):
        super().setup_ui()
        label_func = QLabel("收據轉換功能（還沒寫）")
        self.layout().addWidget(label_func)


class OtherTransferWindow(BaseFuncWindow):
    def __init__(self, title="其他轉換功能（還沒寫）"):
        super().__init__(title)

    def setup_ui(self):
        super().setup_ui()
        label_func = QLabel("其他轉換功能")
        self.layout().addWidget(label_func)


class MainWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("地藏王廟小工具")
        self.setMinimumSize(400, 300)
        self.layout = QVBoxLayout(self)
        self.layout.setAlignment(Qt.AlignmentFlag.AlignCenter)

        self.func_windows = {}

        self.create_buttons()

    def create_buttons(self):
        buttons = [
            ("試算表轉召請文", self.open_excel_to_invitation_window),
            ("試算表轉拔薦文疏-本", self.open_invitation_transfer_window),
            ("試算表轉收據", self.open_receipt_transfer_window),
            ("其他轉換功能", self.open_other_transfer_window),
        ]

        for text, slot in buttons:
            btn = QPushButton(text)
            btn.setFixedSize(200, 40)
            btn.clicked.connect(slot)
            self.layout.addWidget(btn)

    def open_func_window(self, window_class, title):
        if title in self.func_windows:
            win = self.func_windows[title]
            if win.isVisible():
                win.raise_()
                win.activateWindow()
                return
        win = window_class(title)
        win.setAttribute(Qt.WidgetAttribute.WA_DeleteOnClose)
        win.destroyed.connect(lambda: self.func_windows.pop(title, None))
        self.func_windows[title] = win
        win.show()

    def open_excel_to_invitation_window(self):
        self.open_func_window(ExcelToInvitationWindow, "試算表轉召請文")

    def open_invitation_transfer_window(self):
        self.open_func_window(InvitationTransferWindow, "試算表轉拔薦文疏-本")

    def open_receipt_transfer_window(self):
        self.open_func_window(ReceiptTransferWindow, "試算表轉收據")

    def open_other_transfer_window(self):
        self.open_func_window(OtherTransferWindow, "其他轉換功能")

    def closeEvent(self, event):
        for win in list(self.func_windows.values()):
            win.close()
        event.accept()


if __name__ == "__main__":
    import sys
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())

